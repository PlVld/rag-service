import logging
import os
from docx import Document as DocxDocument
from .base import BaseCleaner

logger = logging.getLogger(__name__)


class DOCXCleaner(BaseCleaner):
    """
    Очищает DOCX документ, извлекая текст и преобразуя в Markdown.
    Таблицы конвертируются в Markdown-таблицы.
    """

    def __init__(self, ignore_images: bool = True):
        self.ignore_images = ignore_images

    def clean(self, text: str, **kwargs) -> str:
        """
        Извлекает текст и таблицы из DOCX файла, возвращает Markdown.
        Если передан путь к файлу - обрабатывает файл, если передан текст - возвращает как есть.
        Может принимать как путь к файлу, так и уже загруженный объект Document.

        Args:
            text: Путь к DOCX файлу или текстовое содержимое

        Returns:
            str: Текст в формате Markdown
        """
        if not text:
            return ""

        try:
            # Если text - это путь к существующему DOCX файлу, обрабатываем его
            if isinstance(text, str) and text.endswith('.docx') and os.path.exists(text):
                logger.info(f"Starting DOCX extraction from: {text}")
                doc = DocxDocument(text)
            else:
                # Если уже передано текстовое содержимое (а не путь к файлу), возвращаем как есть
                logger.info("Text content already provided, skipping DOCX processing")
                return str(text)  # Просто возвращаем текст как строку
            text_content = ""

            # Обработка параграфов
            for paragraph in doc.paragraphs:
                # Распознаём заголовки по стилям
                style_name = paragraph.style.name.lower() if paragraph.style else ''
                if 'heading 1' in style_name:
                    text_content += f"# {paragraph.text}\n\n"
                elif 'heading 2' in style_name:
                    text_content += f"## {paragraph.text}\n\n"
                elif 'heading 3' in style_name:
                    text_content += f"### {paragraph.text}\n\n"
                else:
                    text_content += f"{paragraph.text}\n\n"

            # Обработка таблиц
            for table in doc.tables:
                if not table.rows:
                    continue
                markdown_table = self._table_to_markdown(table)
                text_content += markdown_table + "\n\n"

            logger.info(f"Completed DOCX extraction, result length: {len(text_content)}")
            return text_content.strip()

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {text}: {e}", exc_info=True)
            return ""

    @staticmethod
    def _table_to_markdown(table) -> str:
        """Конвертирует таблицу python-docx в Markdown-формат."""
        rows = table.rows
        if not rows:
            return ""

        num_cols = len(rows[0].cells)
        markdown_lines = []

        table_data = []
        for row in rows:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if len(row_data) < num_cols:
                row_data.extend([''] * (num_cols - len(row_data)))
            table_data.append(row_data)

        for i, row in enumerate(table_data):
            markdown_lines.append('| ' + ' | '.join(row) + ' |')
            if i == 0:
                separator = '|' + '|'.join([' --- ' for _ in range(num_cols)]) + '|'
                markdown_lines.append(separator)

        return '\n'.join(markdown_lines)